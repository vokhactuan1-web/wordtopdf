"""
Word to PDF Converter
Hỗ trợ Unicode đầy đủ cho tiếng Việt
"""
from pathlib import Path
from typing import Optional

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY

from ..logging.logger_setup import get_logger
from ..io.file_handler import FileHandler
from .excel_to_pdf import FontManager

logger = get_logger(__name__)


class WordToPDFConverter:
    """Class chuyển đổi Word sang PDF"""
    
    def __init__(self):
        self.font_regular, self.font_bold = FontManager.register_fonts()
        self.styles = self._create_styles()
    
    def _create_styles(self) -> dict:
        """Tạo các style cho PDF"""
        base_styles = getSampleStyleSheet()
        
        styles = {
            'Normal': ParagraphStyle(
                'CustomNormal',
                parent=base_styles['Normal'],
                fontName=self.font_regular,
                fontSize=11,
                leading=14,
                spaceAfter=6,
                alignment=TA_LEFT
            ),
            'Heading1': ParagraphStyle(
                'CustomHeading1',
                parent=base_styles['Heading1'],
                fontName=self.font_bold,
                fontSize=16,
                textColor=colors.HexColor('#2C3E50'),
                spaceAfter=12,
                spaceBefore=12,
                alignment=TA_LEFT
            ),
            'Heading2': ParagraphStyle(
                'CustomHeading2',
                parent=base_styles['Heading2'],
                fontName=self.font_bold,
                fontSize=14,
                textColor=colors.HexColor('#34495E'),
                spaceAfter=10,
                spaceBefore=10,
                alignment=TA_LEFT
            ),
            'Heading3': ParagraphStyle(
                'CustomHeading3',
                parent=base_styles['Heading3'],
                fontName=self.font_bold,
                fontSize=12,
                textColor=colors.HexColor('#7F8C8D'),
                spaceAfter=8,
                spaceBefore=8,
                alignment=TA_LEFT
            ),
            'Bullet': ParagraphStyle(
                'CustomBullet',
                parent=base_styles['Normal'],
                fontName=self.font_regular,
                fontSize=11,
                leftIndent=20,
                bulletIndent=10,
                spaceAfter=6
            )
        }
        
        return styles
    
    def convert(self, input_path: Path, output_path: Optional[Path] = None) -> Path:
        """
        Chuyển Word sang PDF
        
        Args:
            input_path: Đường dẫn file Word
            output_path: Đường dẫn file PDF output (tùy chọn)
            
        Returns:
            Path: Đường dẫn file PDF đã tạo
        """
        output_path = FileHandler.ensure_output_path(output_path, input_path, '.pdf')
        
        logger.info(f"Đang đọc Word: {input_path.name}")
        doc = Document(input_path)
        
        # Tạo PDF document
        pdf_doc = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        elements = self._process_document(doc)
        
        logger.info("Đang tạo PDF...")
        pdf_doc.build(elements)
        logger.info(f"Đã tạo PDF: {output_path}")
        
        return output_path
    
    def _process_document(self, doc: Document) -> list:
        """Xử lý document và tạo elements cho PDF"""
        elements = []
        
        for para in doc.paragraphs:
            element = self._process_paragraph(para)
            if element:
                elements.append(element)
        
        # Xử lý tables
        for table in doc.tables:
            table_element = self._process_table(table)
            if table_element:
                elements.append(table_element)
                elements.append(Spacer(1, 0.2*inch))
        
        logger.info(f"Đã xử lý {len(elements)} elements")
        return elements
    
    def _process_paragraph(self, para) -> Optional[Paragraph]:
        """Xử lý một paragraph"""
        text = para.text.strip()
        if not text:
            return Spacer(1, 0.1*inch)
        
        # Xác định style dựa trên paragraph style
        style_name = para.style.name
        
        if 'Heading 1' in style_name:
            style = self.styles['Heading1']
        elif 'Heading 2' in style_name:
            style = self.styles['Heading2']
        elif 'Heading 3' in style_name:
            style = self.styles['Heading3']
        elif 'List' in style_name or 'Bullet' in style_name:
            text = f"• {text}"
            style = self.styles['Bullet']
        else:
            style = self.styles['Normal']
        
        # Escape HTML special characters
        text = self._escape_html(text)
        
        # Áp dụng formatting (bold, italic)
        text = self._apply_formatting(para, text)
        
        return Paragraph(text, style)
    
    def _escape_html(self, text: str) -> str:
        """Escape HTML special characters"""
        text = text.replace('&', '&amp;')
        text = text.replace('<', '&lt;')
        text = text.replace('>', '&gt;')
        return text
    
    def _apply_formatting(self, para, text: str) -> str:
        """Áp dụng bold/italic formatting"""
        # Kiểm tra runs để áp dụng formatting
        formatted_text = ""
        for run in para.runs:
            run_text = self._escape_html(run.text)
            
            if run.bold and run.italic:
                formatted_text += f"<b><i>{run_text}</i></b>"
            elif run.bold:
                formatted_text += f"<b>{run_text}</b>"
            elif run.italic:
                formatted_text += f"<i>{run_text}</i>"
            else:
                formatted_text += run_text
        
        return formatted_text if formatted_text else text
    
    def _process_table(self, table) -> Optional[Table]:
        """Xử lý bảng từ Word"""
        data = []
        
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                cell_text = self._escape_html(cell_text)
                row_data.append(cell_text)
            data.append(row_data)
        
        if not data:
            return None
        
        # Tạo bảng PDF
        pdf_table = Table(data)
        pdf_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498DB')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), self.font_bold),
            ('FONTNAME', (0, 1), (-1, -1), self.font_regular),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1),
             [colors.white, colors.HexColor('#ECF0F1')]),
        ]))
        
        return pdf_table


# Hàm helper để sử dụng trực tiếp
def convert_word_to_pdf(input_path: Path, output_path: Optional[Path] = None) -> Path:
    """
    Chuyển Word sang PDF
    
    Args:
        input_path: Đường dẫn file Word
        output_path: Đường dẫn file PDF (tùy chọn)
        
    Returns:
        Path: Đường dẫn file PDF
    """
    converter = WordToPDFConverter()
    return converter.convert(input_path, output_path)